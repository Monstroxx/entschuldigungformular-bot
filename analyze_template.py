#!/usr/bin/env python3
"""Script zur Analyse des Entschuldigungsformular-Templates."""

import os
import sys
from pathlib import Path
from docx import Document
import re

def analyze_docx_template(template_path: str):
    """Analysiert ein DOCX-Template und extrahiert Platzhalter."""
    
    print(f"🔍 Analysiere Template: {template_path}")
    
    try:
        # Lade DOCX
        doc = Document(template_path)
        
        print(f"✅ Template geladen: {len(doc.paragraphs)} Paragraphen, {len(doc.tables)} Tabellen")
        
        # Analysiere Paragraphen
        print("\n📝 PARAGRAPHEN:")
        placeholders = set()
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"  {i+1:2d}. {paragraph.text}")
                
                # Suche nach Platzhaltern (verschiedene Formate)
                patterns = [
                    r'\[([A-Z_]+)\]',  # [PLATZHALTER]
                    r'\{([A-Z_]+)\}',  # {PLATZHALTER}
                    r'<([A-Z_]+)>',    # <PLATZHALTER>
                    r'___+',           # ____ (Unterstriche)
                    r'\.\.\.+',        # ... (Punkte)
                ]
                
                for pattern in patterns:
                    matches = re.findall(pattern, paragraph.text)
                    placeholders.update(matches)
        
        # Analysiere Tabellen
        print(f"\n📊 TABELLEN ({len(doc.tables)} gefunden):")
        for table_idx, table in enumerate(doc.tables):
            print(f"  Tabelle {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                        
                        # Suche Platzhalter in Zellen
                        for pattern in patterns:
                            matches = re.findall(pattern, cell_text)
                            placeholders.update(matches)
                
                if row_text:
                    print(f"    Zeile {row_idx + 1}: {' | '.join(row_text)}")
        
        # Zeige gefundene Platzhalter
        print(f"\n🎯 GEFUNDENE PLATZHALTER ({len(placeholders)}):")
        for placeholder in sorted(placeholders):
            print(f"  • {placeholder}")
        
        # Generiere Template-Mapping basierend auf dem verbesserten Formular
        print(f"\n📋 VORGESCHLAGENES TEMPLATE-MAPPING:")
        template_mapping = {
            "VORNAME": "first_name",
            "NACHNAME": "last_name", 
            "AKTUELLES_DATUM": "current_date",
            "BERUFSKOLLEG": "school_name",
            "HIT12": "class_name",
            "2025/2026": "school_year"
        }
        
        for placeholder in sorted(placeholders):
            if placeholder in template_mapping:
                print(f"  {placeholder} -> {template_mapping[placeholder]}")
            else:
                print(f"  {placeholder} -> ??? (NEUER PLATZHALTER)")
        
        # Zeige Formular-Struktur
        print(f"\n📋 FORMULAR-STRUKTUR:")
        print("  • Berufskolleg Bergisch Gladbach HIT12 2025/2026")
        print("  • Entschuldigungsformular (Titel)")
        print("  • Nachname, Vorname")
        print("  • Anrede: Sehr geehrter Herr Bruns,")
        print("  • Einleitung: Ich entschuldige mein Fehlen...")
        print("  • Tabelle: 6 Spalten (leer, leer, 1./2., 3./4., 5./6., 7./8.)")
        print("  • Anmerkung: Klausurtermine müssen gekennzeichnet...")
        print("  • Hinweis: Bei Bedarf die oben stehende Tabelle duplizieren")
        print("  • Ort, Datum")
        print("  • Unterschrift (bei Minderjährigen...)")
        print("  • Schulkonferenz-Beschluss vom 30.09.2024")
        
        return placeholders
        
    except Exception as e:
        print(f"❌ Fehler beim Analysieren: {e}")
        return set()

def export_to_markdown(template_path: str, output_path: str = None):
    """Exportiert DOCX zu Markdown für bessere Analyse."""
    
    if output_path is None:
        output_path = template_path.replace('.docx', '.md')
    
    try:
        # Verwende pandoc falls verfügbar
        import subprocess
        result = subprocess.run([
            'pandoc', template_path, '-o', output_path
        ], capture_output=True, text=True)
        
        if result.returncode == 0:
            print(f"✅ Markdown exportiert: {output_path}")
            return output_path
        else:
            print(f"⚠️ Pandoc nicht verfügbar, verwende Python-Export")
            return None
            
    except FileNotFoundError:
        print("⚠️ Pandoc nicht installiert. Installiere mit: sudo apt install pandoc")
        return None

def main():
    """Hauptfunktion."""
    
    # Suche nach Template-Dateien
    template_dir = Path("formular_examples")
    docx_files = list(template_dir.glob("*.docx"))
    
    if not docx_files:
        print("❌ Keine DOCX-Dateien in formular_examples/ gefunden")
        return
    
    print(f"📁 Gefundene DOCX-Dateien: {len(docx_files)}")
    
    for docx_file in docx_files:
        print(f"\n{'='*60}")
        print(f"📄 Analysiere: {docx_file.name}")
        print('='*60)
        
        # Analysiere DOCX
        placeholders = analyze_docx_template(str(docx_file))
        
        # Exportiere zu Markdown
        md_file = export_to_markdown(str(docx_file))
        if md_file:
            print(f"\n📖 Markdown-Version erstellt: {md_file}")
            print("   Du kannst diese Datei öffnen, um das Template besser zu verstehen.")

if __name__ == "__main__":
    main()
