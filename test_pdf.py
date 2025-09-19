#!/usr/bin/env python3
"""Test Script für PDF-Konvertierung."""

import os
import sys
sys.path.append('.')

from bot.utils.pdf_converter import PDFConverter

def test_pdf_conversion():
    """Teste die PDF-Konvertierung."""
    print("🔍 Teste PDF-Konvertierung...")
    
    # Erstelle eine Test-DOCX-Datei
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test Entschuldigungsformular', 0)
    doc.add_paragraph('Dies ist ein Test-Dokument für die PDF-Konvertierung.')
    
    # Erstelle Tabelle
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Datum'
    hdr_cells[1].text = 'Stunde'
    hdr_cells[2].text = 'Fach'
    hdr_cells[3].text = 'Grund'
    
    # Daten
    row_cells = table.rows[1].cells
    row_cells[0].text = '19.09.2025'
    row_cells[1].text = '1-2'
    row_cells[2].text = 'Mathematik'
    row_cells[3].text = 'Krankheit'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = '19.09.2025'
    row_cells[1].text = '3-4'
    row_cells[2].text = 'Deutsch'
    row_cells[3].text = 'Krankheit'
    
    # Speichere DOCX
    test_docx = 'test_formular.docx'
    doc.save(test_docx)
    print(f"✅ Test-DOCX erstellt: {test_docx}")
    
    # Teste PDF-Konvertierung
    converter = PDFConverter()
    pdf_path = converter.convert_docx_to_pdf(test_docx)
    
    if pdf_path and os.path.exists(pdf_path):
        print(f"✅ PDF erfolgreich erstellt: {pdf_path}")
        print(f"📄 PDF-Größe: {os.path.getsize(pdf_path)} bytes")
        
        # Lösche Test-Dateien
        os.remove(test_docx)
        os.remove(pdf_path)
        print("✅ Test-Dateien gelöscht")
        return True
    else:
        print("❌ PDF-Konvertierung fehlgeschlagen")
        if os.path.exists(test_docx):
            os.remove(test_docx)
        return False

if __name__ == "__main__":
    success = test_pdf_conversion()
    sys.exit(0 if success else 1)
