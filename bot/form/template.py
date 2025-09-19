"""Formular Template für den Entschuldigungsformular Bot."""

import os
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)


class FormTemplate:
    """Template für Entschuldigungsformulare."""
    
    def __init__(self, template_path: str = None):
        """Initialisiert das Formular Template."""
        if template_path is None:
            # Verwende das Beispiel-Formular als Template
            template_path = os.path.join(
                os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
                "formular_examples",
                "2025-Entschuldigungsformular.docx"
            )
        
        self.template_path = template_path
        self.template_exists = os.path.exists(template_path)
    
    def create_base_template(self) -> Document:
        """Erstellt ein Basis-Template basierend auf dem echten Formular."""
        doc = Document()
        
        # Titel
        title = doc.add_heading("Entschuldigungsformular", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Leerzeile
        doc.add_paragraph()
        
        # Name
        name_para = doc.add_paragraph()
        name_para.add_run("Nachname, Vorname: ").bold = True
        name_para.add_run("[NACHNAME], [VORNAME]")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Anrede
        greeting = doc.add_paragraph("Sehr geehrter Herr Bruns,")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Einleitung
        intro = doc.add_paragraph("Ich entschuldige mein Fehlen für die Unterrichtsstunden an folgenden Tagen:")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Anmerkung
        note = doc.add_paragraph("Anmerkung: Klausurtermine müssen gekennzeichnet und mit Attest entschuldigt werden.")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Hinweis
        hint = doc.add_paragraph("Bei Bedarf die oben stehende Tabelle duplizieren.")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Erstelle Tabelle für Fehlzeiten (basierend auf dem echten Format)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Header-Zeile
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '1./2.'
        hdr_cells[1].text = '3./4.'
        hdr_cells[2].text = '5./6.'
        hdr_cells[3].text = '7./8.'
        hdr_cells[4].text = 'Fach'
        hdr_cells[5].text = 'Lehrerkürzel'
        
        # Beispiel-Zeile für Montag
        row_cells = table.add_row().cells
        row_cells[0].text = 'Montag'
        row_cells[1].text = '[DATUM_1]'
        row_cells[2].text = '[STUNDE_1]'
        row_cells[3].text = '[STUNDE_2]'
        row_cells[4].text = '[FACH_1]'
        row_cells[5].text = '[LEHRER_1]'
        
        # Leerzeile
        doc.add_paragraph()
        
        # Ort und Datum
        location_date = doc.add_paragraph()
        location_date.add_run("Ort, Datum: ").bold = True
        location_date.add_run("Bergisch Gladbach, [AKTUELLES_DATUM]")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Unterschrift
        signature = doc.add_paragraph()
        signature.add_run("Unterschrift")
        signature.add_run("\n(bei Minderjährigen von einem Erziehungsberechtigten)")
        
        return doc
    
    def load_template(self) -> Document:
        """Lädt das Formular Template."""
        if self.template_exists:
            try:
                return Document(self.template_path)
            except Exception as e:
                logger.warning(f"Fehler beim Laden des Templates, verwende Basis-Template: {e}")
                return self.create_base_template()
        else:
            logger.info("Template nicht gefunden, erstelle Basis-Template")
            return self.create_base_template()
    
    def fill_template(self, data: Dict[str, Any]) -> Document:
        """Füllt das Template mit den übergebenen Daten aus."""
        doc = self.load_template()
        
        # Ersetze Platzhalter basierend auf dem echten Formular
        replacements = {
            "[VORNAME]": data.get("first_name", ""),
            "[NACHNAME]": data.get("last_name", ""),
            "[AKTUELLES_DATUM]": data.get("current_date", datetime.now().strftime("%d.%m.%Y")),
            "[DATUM_1]": data.get("start_date", ""),
            "[STUNDE_1]": data.get("hour_1", ""),
            "[STUNDE_2]": data.get("hour_2", ""),
            "[FACH_1]": data.get("subject_1", ""),
            "[LEHRER_1]": data.get("teacher_1", ""),
        }
        
        # Ersetze in allen Paragraphen
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Ersetze in Tabellen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Füge Stundenplan hinzu falls vorhanden
        if "schedule" in data and data["schedule"]:
            self._add_schedule_table(doc, data["schedule"], data.get("absence_periods", []))
        
        return doc
    
    def _add_schedule_table(self, doc: Document, schedule: List[Dict[str, str]], absence_periods: List[Dict[str, Any]]):
        """Fügt eine detaillierte Stundenplan-Tabelle basierend auf dem echten Format hinzu."""
        
        # Finde die bestehende Tabelle
        tables = doc.tables
        if not tables:
            return
        
        table = tables[0]  # Verwende die erste Tabelle
        
        # Erstelle eine neue Tabelle für jede Fehlzeit
        for period in absence_periods:
            start_date = period["start"]
            end_date = period["end"]
            
            # Erstelle Datum-Liste
            current_date = start_date
            while current_date <= end_date:
                # Füge Zeile für jeden Tag hinzu
                row_cells = table.add_row().cells
                
                # Wochentag
                weekday = current_date.strftime("%A")
                weekday_de = {
                    "Monday": "Montag",
                    "Tuesday": "Dienstag", 
                    "Wednesday": "Mittwoch",
                    "Thursday": "Donnerstag",
                    "Friday": "Freitag"
                }.get(weekday, weekday)
                
                row_cells[0].text = weekday_de
                row_cells[1].text = current_date.strftime("%d.%m.%Y")
                
                # Stunden und Fächer basierend auf Stundenplan
                if schedule:
                    # Erste Stunde
                    first_hour = schedule[0] if len(schedule) > 0 else {"hour": "", "subject": ""}
                    row_cells[2].text = first_hour["hour"]
                    row_cells[3].text = first_hour["subject"]
                    
                    # Zweite Stunde (falls vorhanden)
                    if len(schedule) > 1:
                        second_hour = schedule[1]
                        row_cells[4].text = second_hour["hour"]
                        row_cells[5].text = second_hour["subject"]
                    else:
                        row_cells[4].text = ""
                        row_cells[5].text = ""
                else:
                    row_cells[2].text = ""
                    row_cells[3].text = ""
                    row_cells[4].text = ""
                    row_cells[5].text = ""
                
                current_date += datetime.timedelta(days=1)
    
    def save_document(self, doc: Document, filename: str = None) -> str:
        """Speichert das Dokument und gibt den Pfad zurück."""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"entschuldigung_{timestamp}.docx"
        
        # Erstelle output Ordner falls nicht vorhanden
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "output")
        os.makedirs(output_dir, exist_ok=True)
        
        file_path = os.path.join(output_dir, filename)
        doc.save(file_path)
        
        return file_path
