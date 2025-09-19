"""Template für das echte Entschuldigungsformular."""

import os
from datetime import datetime, timedelta
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)


class RealFormTemplate:
    """Template basierend auf dem echten Entschuldigungsformular."""
    
    def __init__(self, template_path: str = None):
        """Initialisiert das Template."""
        if template_path is None:
            # Prüfe verschiedene mögliche Pfade
            possible_paths = [
                # Railway/Production Pfade
                os.path.join("/app", "templates", "2025-Entschuldigungsformular.docx"),
                os.path.join("/app", "templates", "entschuldigung_template.docx"),
                # Lokale Entwicklung Pfade
                os.path.join(
                    os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
                    "templates",
                    "2025-Entschuldigungsformular.docx"
                ),
                os.path.join(
                    os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
                    "templates",
                    "entschuldigung_template.docx"
                ),
                os.path.join(
                    os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
                    "formular_examples",
                    "2025-Entschuldigungsformular.docx"
                )
            ]
            
            # Finde den ersten existierenden Pfad
            template_path = None
            for path in possible_paths:
                if os.path.exists(path):
                    template_path = path
                    break
            
            # Fallback auf None, wenn nichts gefunden wird
            if template_path is None:
                template_path = possible_paths[0]  # Verwende den ersten Pfad als Standard
        
        self.template_path = template_path
        self.template_exists = os.path.exists(template_path)
        
        # Debug-Logging
        logger.info(f"Template-Pfad: {template_path}")
        logger.info(f"Template existiert: {self.template_exists}")
        if not self.template_exists:
            logger.warning(f"Template nicht gefunden unter: {template_path}")
            # Liste verfügbare Verzeichnisse
            base_dir = os.path.dirname(template_path)
            if os.path.exists(base_dir):
                logger.info(f"Verfügbare Dateien in {base_dir}: {os.listdir(base_dir)}")
            else:
                logger.warning(f"Verzeichnis existiert nicht: {base_dir}")
    
    def load_template(self) -> Document:
        """Lädt das echte Formular-Template."""
        if self.template_exists:
            try:
                return Document(self.template_path)
            except Exception as e:
                logger.error(f"Fehler beim Laden des echten Templates: {e}")
                return self.create_fallback_template()
        else:
            logger.warning("Echtes Template nicht gefunden, verwende Fallback")
            return self.create_fallback_template()
    
    def create_fallback_template(self) -> Document:
        """Erstellt ein Fallback-Template basierend auf dem verbesserten Formular."""
        doc = Document()
        
        # Schulname (zentriert)
        school_name = doc.add_paragraph()
        school_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        school_name.add_run("Berufskolleg Bergisch Gladbach HIT12 2025/2026")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Titel
        title = doc.add_paragraph()
        title.add_run("Entschuldigungsformular").bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Leerzeile
        doc.add_paragraph()
        
        # Name
        name_para = doc.add_paragraph()
        name_para.add_run("Nachname, Vorname: ").bold = True
        name_para.add_run("[NACHNAME], [VORNAME]")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Anrede
        greeting = doc.add_paragraph("Sehr geehrter Herr Bruns,")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Einleitung
        intro = doc.add_paragraph("Ich entschuldige mein Fehlen für die Unterrichtsstunden an folgenden Tagen:")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Erstelle Tabelle im verbesserten Format (6 Spalten)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Header-Zeile (leer für die ersten beiden Spalten)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''  # Leer
        hdr_cells[1].text = ''  # Leer
        hdr_cells[2].text = '1./2.'
        hdr_cells[3].text = '3./4.'
        hdr_cells[4].text = '5./6.'
        hdr_cells[5].text = '7./8.'
        
        # Leerzeile
        doc.add_paragraph()
        
        # Anmerkung
        note = doc.add_paragraph("Anmerkung: Klausurtermine müssen gekennzeichnet und mit Attest entschuldigt werden.")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Hinweis
        hint = doc.add_paragraph("Bei Bedarf die oben stehende Tabelle duplizieren.")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Ort und Datum
        location_date = doc.add_paragraph()
        location_date.add_run("Ort, Datum: ").bold = True
        location_date.add_run("Bergisch Gladbach, [AKTUELLES_DATUM]")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Unterschrift
        signature = doc.add_paragraph()
        signature.add_run("Unterschrift")
        signature.add_run("\n(bei Minderjährigen von einem Erziehungsberechtigten)")
        
        # Leerzeile
        doc.add_paragraph()
        
        # Schulkonferenz-Beschluss
        note = doc.add_paragraph()
        note.add_run("Beschluss der Schulkonferenz vom 30.09.2024").bold = True
        note.add_run("\nMinderjährigen mit Unterschrift der Erziehungsberechtigten) oder ein Attest in Papierform unaufgefordert bei der Klassenleitung eingereicht worden sein. Später eingereichte Entschuldigungen werden nicht mehr akzeptiert und führen zu unentschuldigten Fehlzeiten auf dem Zeugnis.")
        
        return doc
    
    def fill_template(self, data: Dict[str, Any]) -> Document:
        """Füllt das Template mit den übergebenen Daten aus."""
        doc = self.load_template()
        
        # Ersetze Platzhalter (für Fallback-Template)
        replacements = {
            "[VORNAME]": data.get("first_name", ""),
            "[NACHNAME]": data.get("last_name", ""),
            "[AKTUELLES_DATUM]": data.get("current_date", datetime.now().strftime("%d.%m.%Y")),
        }
        
        # Ersetze in allen Paragraphen
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Ersetze in Tabellen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Für das echte Template: Suche nach spezifischen Textmustern und ersetze sie
        if self.template_exists:
            self._fill_real_template(doc, data)
        
        # Füge Fehlzeiten-Tabelle hinzu
        if "absence_periods" in data and data["absence_periods"]:
            self._add_absence_table(doc, data["absence_periods"], data.get("schedule", []))
        
        return doc
    
    def _fill_real_template(self, doc: Document, data: Dict[str, Any]) -> None:
        """Füllt das echte Template mit den übergebenen Daten aus."""
        first_name = data.get("first_name", "")
        last_name = data.get("last_name", "")
        reason = data.get("reason", "")
        current_date = data.get("current_date", datetime.now().strftime("%d.%m.%Y"))
        
        # Suche nach Name-Feldern und ersetze sie
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # Suche nach "Nachname, Vorname:" und ersetze den Inhalt danach
            if "Nachname, Vorname:" in text:
                # Finde den Text nach "Nachname, Vorname:"
                parts = text.split("Nachname, Vorname:")
                if len(parts) > 1:
                    # Ersetze alles nach "Nachname, Vorname:" mit dem echten Namen
                    paragraph.text = f"Nachname, Vorname: {last_name}, {first_name}"
            
            # Suche nach Grund/Reason und ersetze
            if "Grund:" in text and not reason in text:
                parts = text.split("Grund:")
                if len(parts) > 1:
                    paragraph.text = f"Grund: {reason}"
            
            # Suche nach Datum und ersetze
            if "Datum:" in text and not current_date in text:
                parts = text.split("Datum:")
                if len(parts) > 1:
                    paragraph.text = f"Datum: {current_date}"
        
        # Suche in Tabellen nach Platzhaltern
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        
                        # Ersetze Name-Platzhalter
                        if "[NACHNAME]" in text:
                            paragraph.text = text.replace("[NACHNAME]", last_name)
                        if "[VORNAME]" in text:
                            paragraph.text = text.replace("[VORNAME]", first_name)
                        if "[GRUND]" in text:
                            paragraph.text = text.replace("[GRUND]", reason)
                        if "[DATUM]" in text:
                            paragraph.text = text.replace("[DATUM]", current_date)
    
    def _add_absence_table(self, doc: Document, absence_periods: List[Dict[str, Any]], schedule: List[Dict[str, str]]):
        """Fügt die Fehlzeiten-Tabelle hinzu."""
        
        # Finde die bestehende Tabelle
        tables = doc.tables
        if not tables:
            return
        
        table = tables[0]  # Verwende die erste Tabelle
        
        # Entferne die Beispiel-Zeile falls vorhanden
        if len(table.rows) > 1:
            # Prüfe ob es eine Beispiel-Zeile ist
            first_data_row = table.rows[1]
            if "Montag" in first_data_row.cells[0].text:
                # Entferne Beispiel-Zeilen
                while len(table.rows) > 1:
                    table._element.remove(table.rows[1]._element)
        
        # Füge Zeilen für jede Fehlzeit hinzu
        for period in absence_periods:
            start_date = period["start"]
            end_date = period["end"]
            
            # Erstelle Datum-Liste
            current_date = start_date
            while current_date <= end_date:
                # Füge Zeile für jeden Tag hinzu
                row_cells = table.add_row().cells
                
                # Wochentag und Datum in die ersten beiden Spalten
                weekday = current_date.strftime("%A")
                weekday_de = {
                    "Monday": "Montag",
                    "Tuesday": "Dienstag", 
                    "Wednesday": "Mittwoch",
                    "Thursday": "Donnerstag",
                    "Friday": "Freitag"
                }.get(weekday, weekday)
                
                row_cells[0].text = weekday_de
                row_cells[1].text = current_date.strftime("%d.%m.%Y")
                
                # Stunden-Spalten (3./4., 5./6., 7./8.) - basierend auf Stundenplan
                if schedule:
                    # Erste Stunde (1./2.)
                    first_hour = schedule[0] if len(schedule) > 0 else {"hour": "", "subject": ""}
                    row_cells[2].text = first_hour["hour"] if first_hour["hour"] else ""
                    
                    # Zweite Stunde (3./4.)
                    if len(schedule) > 1:
                        second_hour = schedule[1]
                        row_cells[3].text = second_hour["hour"] if second_hour["hour"] else ""
                    else:
                        row_cells[3].text = ""
                    
                    # Dritte Stunde (5./6.)
                    if len(schedule) > 2:
                        third_hour = schedule[2]
                        row_cells[4].text = third_hour["hour"] if third_hour["hour"] else ""
                    else:
                        row_cells[4].text = ""
                    
                    # Vierte Stunde (7./8.)
                    if len(schedule) > 3:
                        fourth_hour = schedule[3]
                        row_cells[5].text = fourth_hour["hour"] if fourth_hour["hour"] else ""
                    else:
                        row_cells[5].text = ""
                else:
                    # Leer lassen wenn kein Stundenplan
                    row_cells[2].text = ""
                    row_cells[3].text = ""
                    row_cells[4].text = ""
                    row_cells[5].text = ""
                
                current_date += timedelta(days=1)
    
    def save_document(self, doc: Document, filename: str = None) -> str:
        """Speichert das Dokument und gibt den Pfad zurück."""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"entschuldigung_{timestamp}.docx"
        
        # Erstelle output Ordner falls nicht vorhanden
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "output")
        os.makedirs(output_dir, exist_ok=True)
        
        file_path = os.path.join(output_dir, filename)
        doc.save(file_path)
        
        return file_path
