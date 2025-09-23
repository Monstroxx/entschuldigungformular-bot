#!/usr/bin/env python3
"""
Python Script für Template-Manipulation und Tabellen-Generierung
Wird vom TypeScript Bot aufgerufen
"""

import sys
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_template(template_path):
    """Lädt das DOCX Template"""
    try:
        return Document(template_path)
    except Exception as e:
        print(f"Fehler beim Laden des Templates: {e}", file=sys.stderr)
        return None

def replace_placeholders(doc, data):
    """Ersetzt Platzhalter im Dokument"""
    # Ersetze alle Platzhalter im gesamten Dokument
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            if '[NACHNAME]' in text:
                text = text.replace('[NACHNAME]', data.get('lastName', ''))
            if '[VORNAME]' in text:
                text = text.replace('[VORNAME]', data.get('firstName', ''))
            if '[GRUND]' in text:
                text = text.replace('[GRUND]', data.get('reason', ''))
            if '[ORT]' in text:
                text = text.replace('[ORT]', data.get('location', 'Bergisch Gladbach'))
            if '[DATUM]' in text:
                text = text.replace('[DATUM]', data.get('currentDate', ''))
            if '[LEHRER]' in text:
                text = text.replace('[LEHRER]', data.get('teacherLastName', 'Müller'))
            run.text = text

def find_existing_table(doc):
    """Findet die bestehende Tabelle im Template"""
    for table in doc.tables:
        # Suche nach einer Tabelle mit den typischen Header-Zellen
        if len(table.rows) > 0 and len(table.rows[0].cells) >= 6:
            first_row = table.rows[0]
            header_text = ' '.join([cell.text.strip() for cell in first_row.cells])
            if '1./2.' in header_text and '3./4.' in header_text:
                return table
    return None

def manipulate_existing_table(doc, data):
    """Manipuliert die bestehende Tabelle im Template"""
    table = find_existing_table(doc)
    if not table:
        print("Warnung: Keine bestehende Tabelle gefunden", file=sys.stderr)
        return

    # Lösche alle Zeilen außer der Header-Zeile
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)

    # Füge Fehlzeiten hinzu
    absence_periods = data.get('absencePeriods', [])
    schedule = data.get('schedule', [])

    # Gruppiere Fehlzeiten nach Wochen
    weeks = group_absence_periods_by_week(absence_periods)

    for week in weeks:
        for period in week:
            start_date = datetime.fromisoformat(period['start'].replace('Z', '+00:00'))
            end_date = datetime.fromisoformat(period['end'].replace('Z', '+00:00'))
            
            current_date = start_date
            while current_date <= end_date:
                # Überspringe Wochenenden
                if current_date.weekday() < 5:  # 0-4 = Montag-Freitag
                    weekday = get_weekday_name(current_date.weekday())
                    date_str = current_date.strftime('%d.%m.%Y')
                    
                    # Hole Stundenplan für diesen Tag
                    schedule_entries = get_schedule_for_day(schedule, weekday)
                    
                    # Füge Zeile hinzu
                    row = table.add_row()
                    row.cells[0].text = weekday
                    row.cells[1].text = date_str
                    row.cells[2].text = schedule_entries.get('1./2.', '')
                    row.cells[3].text = schedule_entries.get('3./4.', '')
                    row.cells[4].text = schedule_entries.get('5./6.', '')
                    row.cells[5].text = schedule_entries.get('7./8.', '')
                
                current_date = current_date.replace(day=current_date.day + 1)

def group_absence_periods_by_week(periods):
    """Gruppiert Fehlzeiten nach Wochen"""
    weeks = []
    current_week = []
    
    for period in periods:
        start_date = datetime.fromisoformat(period['start'].replace('Z', '+00:00'))
        end_date = datetime.fromisoformat(period['end'].replace('Z', '+00:00'))
        
        # Generiere alle Daten im Zeitraum
        current_date = start_date
        while current_date <= end_date:
            # Überspringe Wochenenden
            if current_date.weekday() < 5:  # 0-4 = Montag-Freitag
                current_week.append({
                    'start': current_date.isoformat(),
                    'end': current_date.isoformat()
                })
                
                # Wenn wir 5 Tage haben (Montag-Freitag), starte eine neue Woche
                if len(current_week) >= 5:
                    weeks.append(current_week)
                    current_week = []
            
            current_date = current_date.replace(day=current_date.day + 1)
    
    # Füge verbleibende Tage hinzu
    if current_week:
        weeks.append(current_week)
    
    return weeks

def get_weekday_name(weekday):
    """Konvertiert Wochentag-Nummer zu Name"""
    days = ['Montag', 'Dienstag', 'Mittwoch', 'Donnerstag', 'Freitag']
    return days[weekday]

def get_schedule_for_day(schedule, weekday):
    """Holt Stundenplan für einen bestimmten Wochentag"""
    schedule_entries = {}
    for entry in schedule:
        if entry.get('weekday') == weekday:
            hour = entry.get('hour', '')
            subject = entry.get('subject', '')
            
            # Mappe Stunden zu Zeitblöcken
            if '1.' in hour and '2.' in hour:
                schedule_entries['1./2.'] = subject
            elif '3.' in hour and '4.' in hour:
                schedule_entries['3./4.'] = subject
            elif '5.' in hour and '6.' in hour:
                schedule_entries['5./6.'] = subject
            elif '7.' in hour and '8.' in hour:
                schedule_entries['7./8.'] = subject
    
    return schedule_entries

def process_template(input_path, output_path, data):
    """Hauptfunktion für Template-Verarbeitung"""
    try:
        # Lade Template
        doc = load_template(input_path)
        if not doc:
            return False

        # Ersetze Platzhalter
        replace_placeholders(doc, data)
        
        # Manipuliere bestehende Tabelle
        manipulate_existing_table(doc, data)

        # Speichere Dokument
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"Fehler bei Template-Verarbeitung: {e}", file=sys.stderr)
        return False

def main():
    """Hauptfunktion"""
    if len(sys.argv) != 4:
        print("Usage: python template_processor.py <input_path> <output_path> <data_json>", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    data_json = sys.argv[3]

    try:
        data = json.loads(data_json)
    except json.JSONDecodeError as e:
        print(f"Fehler beim Parsen der JSON-Daten: {e}", file=sys.stderr)
        sys.exit(1)

    success = process_template(input_path, output_path, data)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
